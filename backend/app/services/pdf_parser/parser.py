"""
Document parser that handles PDF (text-based and scanned) and Word files.

Uses PyMuPDF for PDF text extraction, PaddleOCR as a fallback for scanned/image
PDFs, and python-docx for Word documents.

Exports:
    parse(file_path: str) -> ParsedBidDocument
"""

from __future__ import annotations

import logging
import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF

from app.services.pdf_parser.schemas import ParsedBidDocument

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# PaddleOCR lazy import (heavy dependency — only loaded when needed)
# ---------------------------------------------------------------------------
_PADDLE_OCR: Optional[object] = None


def _get_paddle_ocr() -> object:
    """
    Lazily initialise PaddleOCR.

    PaddleOCR is only imported when we actually encounter a scanned PDF
    so that normal text-based PDFs do not pay the import cost.
    """
    global _PADDLE_OCR
    if _PADDLE_OCR is None:
        try:
            from paddleocr import PaddleOCR  # type: ignore[import-untyped]

            _PADDLE_OCR = PaddleOCR(
                lang="ch",
                use_angle_cls=True,
                show_log=False,
                det_db_thresh=0.1,
                det_db_box_thresh=0.2,
                det_limit_side_len=2000,
            )
            logger.info("PaddleOCR initialised successfully.")
        except ImportError:
            logger.warning(
                "PaddleOCR is not installed. Scanned PDFs will not be OCR-processed. "
                "Install with: pip install paddleocr"
            )
            _PADDLE_OCR = False  # type: ignore[assignment]
    return _PADDLE_OCR


# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------


def _extract_text_from_pdf(file_path: str) -> tuple[str, list[str], int]:
    """
    Extract text from a PDF using PyMuPDF.

    Returns:
        (full_text, page_texts, page_count)

    If the amount of extracted text is suspiciously low (indicating a
    scanned/image-only PDF), falls back to PaddleOCR.
    """
    full_text_parts: list[str] = []
    page_texts: list[str] = []

    doc = fitz.open(file_path)
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc.load_page(page_num)
        text = page.get_text()
        page_texts.append(text)
        full_text_parts.append(text)

    doc.close()

    full_text = "\f".join(full_text_parts)

    # Heuristic: if average chars per page is very low, treat as scanned PDF
    total_chars = sum(len(t) for t in page_texts)
    avg_chars = total_chars / max(page_count, 1)

    if avg_chars < 50 and page_count > 0:
        logger.info(
            "Low text density detected (avg %.0f chars/page). Falling back to PaddleOCR.",
            avg_chars,
        )
        return _ocr_pdf(file_path)

    return full_text, page_texts, page_count


def _ocr_pdf(file_path: str) -> tuple[str, list[str], int]:
    """
    OCR a scanned/image PDF page by page using PaddleOCR.

    Each page is rendered to an image, then OCR'd.
    """
    ocr = _get_paddle_ocr()
    if ocr is False:
        # PaddleOCR not available — return what text PyMuPDF could extract
        logger.warning("PaddleOCR unavailable; returning minimal text from PyMuPDF.")
        doc = fitz.open(file_path)
        page_count = len(doc)
        page_texts = []
        for i in range(page_count):
            t = doc.load_page(i).get_text()
            page_texts.append(t)
        doc.close()
        return "\f".join(page_texts), page_texts, page_count

    from paddleocr import PaddleOCR  # type: ignore[import-untyped]
    from PIL import Image, ImageOps
    import numpy as np

    doc = fitz.open(file_path)
    page_count = len(doc)
    page_texts: list[str] = []

    for page_num in range(page_count):
        page = doc.load_page(page_num)
        # Render at high DPI (500) for small-text Chinese documents
        pix = page.get_pixmap(dpi=500)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # Preprocess: only invert if dark-background scan; keep grayscale
        gray = img.convert("L")
        pixels = np.array(gray)
        mean_val = pixels.mean()
        if mean_val < 128:
            gray = ImageOps.invert(gray)

        # For very tall images, split into horizontal strips to stay within
        # OCR model input limits
        height = gray.height
        strip_h = 2500
        all_lines: list[str] = []

        for y in range(0, height, strip_h):
            strip = gray.crop((0, y, gray.width, min(y + strip_h, height)))
            # Skip nearly-empty strips
            if np.array(strip).mean() > 250:
                continue

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                strip.save(tmp, format="PNG")
                tmp_path = tmp.name

            try:
                result = ocr.ocr(tmp_path, cls=True)
                if result and result[0]:
                    for line_info in result[0]:
                        text = line_info[1][0]
                        if text:
                            all_lines.append(text)
            except Exception:
                logger.exception("OCR failed for page %d strip y=%d", page_num + 1, y)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        page_text = "\n".join(all_lines)
        page_texts.append(page_text)

    doc.close()
    full_text = "\f".join(page_texts)
    return full_text, page_texts, page_count


def _extract_text_from_docx(file_path: str) -> tuple[str, list[str], int]:
    """
    Extract text from a Word (.docx) file.

    Since .docx files don't have a native concept of pages, we approximate
    page count at roughly 3000 characters per page.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to parse Word documents. "
            "Install with: pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        paragraphs.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts: list[str] = []
            for cell in row.cells:
                row_texts.append(cell.text)
            paragraphs.append(" | ".join(row_texts))

    full_text = "\n".join(paragraphs)

    # Approximate page count: ~3000 chars per page (Chinese text)
    approx_pages = max(1, len(full_text) // 3000)
    # For docx we return a single "page" containing all text
    page_texts = [full_text]

    return full_text, page_texts, approx_pages


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def parse(file_path: str) -> ParsedBidDocument:
    """
    Parse a bid/tender document and return structured text.

    Handles PDF (including scanned/image-only via PaddleOCR fallback)
    and Word (.docx) files.

    Args:
        file_path: Path to the document file.

    Returns:
        ParsedBidDocument with full_text, page_count populated.
        The remaining fields (score_items, reject_clauses, etc.) are
        left empty — they are filled later by the LLM extractor.

    Raises:
        ValueError: If the file format is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    logger.info("Parsing document: %s (%s)", path.name, ext)

    if ext == ".pdf":
        full_text, page_texts, page_count = _extract_text_from_pdf(str(path))
    elif ext in (".docx", ".doc"):
        if ext == ".doc":
            logger.warning(
                "Legacy .doc format detected. Only .docx is directly supported. "
                "Attempting docx parser; results may be incomplete."
            )
        full_text, page_texts, page_count = _extract_text_from_docx(str(path))
    elif ext == ".txt":
        full_text = path.read_text(encoding="utf-8")
        page_texts = [full_text]
        page_count = 1

    logger.info(
        "Parsing complete — pages=%d chars=%d",
        page_count,
        len(full_text),
    )

    return ParsedBidDocument(
        full_text=full_text,
        page_count=page_count,
    )
