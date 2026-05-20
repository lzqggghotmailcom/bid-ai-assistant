"""Cover-page builder.

Layout (all centred):
  - "投标文件" in 方正小标宋简体 (fallback 黑体), 二号 (22pt), bold
  - Project name in 宋体, 小二 (18pt)
  - Bottom: company name + date, 宋体, 三号 (16pt)
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..styles import (
    FONT_BODY,
    FONT_FZXBS,
    FONT_SIZE_COVER_TITLE,
    FONT_SIZE_COVER_SUBTITLE,
    FONT_SIZE_COVER_BOTTOM,
    LINE_SPACING,
    set_run_font,
    add_paragraph_with_style,
)


def build_cover(doc, bid_info: dict) -> None:
    """Build the cover page.

    Args:
        doc: python-docx Document object.
        bid_info: dict with keys: project_name, company_name, bid_date.
    """
    project_name = bid_info.get('project_name', '')
    company_name = bid_info.get('company_name', '')
    bid_date = bid_info.get('bid_date', '')

    # Top spacer: push title down to roughly one-third of the page
    for _ in range(8):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = LINE_SPACING

    # ---- Title: 投标文件 ----
    add_paragraph_with_style(
        doc, '投标文件',
        font_name=FONT_FZXBS, font_size=FONT_SIZE_COVER_TITLE, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Space between title and subtitle
    for _ in range(2):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = LINE_SPACING

    # ---- Subtitle: Project Name ----
    add_paragraph_with_style(
        doc, project_name,
        font_name=FONT_BODY, font_size=FONT_SIZE_COVER_SUBTITLE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ---- Bottom area: company name + date ----
    # Push to bottom with spacing
    for _ in range(8):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = LINE_SPACING

    add_paragraph_with_style(
        doc, company_name,
        font_name=FONT_BODY, font_size=FONT_SIZE_COVER_BOTTOM,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    date_text = bid_date if bid_date else '年    月    日'
    add_paragraph_with_style(
        doc, date_text,
        font_name=FONT_BODY, font_size=FONT_SIZE_COVER_BOTTOM,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
