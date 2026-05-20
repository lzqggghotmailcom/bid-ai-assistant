"""Table-of-Contents builder.

Inserts a Word TOC field that includes headings levels 1-3.
The TOC is auto-generated when the document is opened in Word / WPS and
the user right-clicks to update fields.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..styles import (
    FONT_TITLE,
    FONT_SIZE_H1,
    LINE_SPACING,
    add_field,
    add_paragraph_with_style,
)


def build_toc(doc) -> None:
    """Add a table-of-contents page to the document.

    Places a "目录" heading followed by a TOC field that picks up
    Outline Level 1-3 paragraphs.
    """
    # ---- TOC Heading ----
    add_paragraph_with_style(
        doc, '目    录',
        font_name=FONT_TITLE, font_size=FONT_SIZE_H1, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Small gap
    gap = doc.add_paragraph()
    gap.paragraph_format.line_spacing = LINE_SPACING

    # ---- TOC Field ----
    # TOC \o "1-3" means include outline levels 1 through 3
    # \h makes entries hyperlinks, \z hides tab-leader in web layout, \u uses applied paragraph outline level
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.line_spacing = 1.0
    add_field(toc_para, ' TOC \\o "1-3" \\h \\z \\u ', '（右键单击此处更新目录）')
