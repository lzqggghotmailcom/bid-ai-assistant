"""Signature / seal page builder.

Separate page with a table-form layout:
  投标人名称（盖章）           | [company name]
  法定代表人或授权代表（签字）  | [blank for signature]
  日期                        | [bid_date or blank]
  联系方式                    | [contact or blank]
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from ..styles import (
    FONT_BODY,
    FONT_SIZE_BODY,
    COLOR_BORDER,
    LINE_SPACING,
    set_run_font,
    set_table_borders,
    add_paragraph_with_style,
)


def build_signature(doc, bid_info: dict) -> None:
    """Add a signature/seal page to the document.

    Args:
        doc: python-docx Document.
        bid_info: dict with keys: company_name, bid_date, contact (optional).
    """
    company = bid_info.get('company_name', '')
    bid_date = bid_info.get('bid_date', '')
    contact = bid_info.get('contact', '')

    # ---- Section heading ----
    add_paragraph_with_style(
        doc, '签署盖章页',
        font_name=FONT_BODY, font_size=FONT_SIZE_BODY, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = LINE_SPACING

    # ---- Signature table ----
    labels = [
        '投标人名称（盖章）',
        '法定代表人或授权代表（签字）',
        '日    期',
        '联系方式',
    ]
    values = [
        company,
        '',          # blank for manual signature
        bid_date or '',
        contact or '',
    ]

    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set column widths: labels ~5cm, values ~10cm
    col_label_width = Cm(5.5)
    col_value_width = Cm(10.5)
    label_col_idx = 0
    value_col_idx = 1

    for row_idx, (label, value) in enumerate(zip(labels, values)):
        row = table.rows[row_idx]
        row.height = Cm(1.2)

        # Label cell
        label_cell = row.cells[label_col_idx]
        label_cell.text = ''
        label_para = label_cell.paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_para.add_run(label)
        set_run_font(label_run, FONT_BODY, FONT_SIZE_BODY, bold=True)
        label_cell.width = col_label_width

        # Value cell
        value_cell = row.cells[value_col_idx]
        value_cell.text = ''
        value_para = value_cell.paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.add_run(value)
        set_run_font(value_run, FONT_BODY, FONT_SIZE_BODY)
        value_cell.width = col_value_width

    # ---- Borders ----
    set_table_borders(table, size=4, color=COLOR_BORDER)

    # ---- Signing instructions ----
    note_spacer = doc.add_paragraph()
    note_spacer.paragraph_format.line_spacing = LINE_SPACING

    note = add_paragraph_with_style(
        doc,
        '注：请在此处加盖单位公章并由法定代表人或授权代表签字。',
        font_name=FONT_BODY, font_size=Pt(9),
    )
    note.paragraph_format.first_line_indent = FONT_SIZE_BODY * 2
