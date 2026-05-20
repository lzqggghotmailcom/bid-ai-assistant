"""Body-content builder.

Renders outline sections with proper heading levels and body text.

Heading specs:
  H1: 黑体 三号 加粗, 段前12磅 段后6磅, numbering 一、二、三
  H2: 黑体 四号 加粗, 段前6磅 段后3磅, numbering （一）（二）（三）
  H3: 楷体 小四 加粗, numbering 1. 2. 3.
  Body: 宋体 小四, 1.5x line-spacing, first-line indent 2 chars
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..styles import (
    FONT_BODY,
    FONT_TITLE,
    FONT_SUBTITLE,
    FONT_SIZE_BODY,
    FONT_SIZE_H1,
    FONT_SIZE_H2,
    FONT_SIZE_H3,
    LINE_SPACING,
    PARA_SPACING_H1_BEFORE,
    PARA_SPACING_H1_AFTER,
    PARA_SPACING_H2_BEFORE,
    PARA_SPACING_H2_AFTER,
    to_chinese_numeral,
    to_bracketed_chinese,
    set_run_font,
    set_paragraph_spacing,
    set_first_line_indent,
    add_paragraph_with_style,
)


def build_content(doc, sections: list[dict]) -> None:
    """Build the body content from the outline sections.

    Args:
        doc: python-docx Document.
        sections: list of section dicts, each with:
            title (str), content (str|None), weight (str|None),
            children (list[dict]|None) - optional nested subsections.
    """
    h1_counter = 0
    _render_sections(doc, sections, level=0,
                     counters={'h1': 0, 'h2': 0, 'h3': 0})


def _render_sections(doc, sections: list[dict], level: int, counters: dict) -> None:
    """Recursively render sections at the given level.

    level 0 = H1, level 1 = H2, level 2 = H3.
    """
    if not sections:
        return

    child_key_map = {0: 'h2', 1: 'h3', 2: None}   # reset child counters
    counter_keys = {0: 'h1', 1: 'h2', 2: 'h3'}

    for section in sections:
        if not isinstance(section, dict):
            continue

        title = section.get('title', '')
        content = section.get('content') or ''
        children = section.get('children')

        # Normalise: if content is None, treat as empty
        if content is None:
            content = ''

        key = counter_keys.get(level)
        if key:
            counters[key] += 1

        # Render heading
        if level == 0:
            _render_h1(doc, title, counters['h1'])
        elif level == 1:
            _render_h2(doc, title, counters['h2'])
        elif level == 2:
            _render_h3(doc, title, counters['h3'])

        # Render body content
        if content:
            _render_body_paragraphs(doc, content)

        # Render children recursively and then reset the deeper counter(s)
        if children and level < 2:
            # Reset the counters for levels deeper than current
            _reset_counters_for_level(counters, level)
            _render_sections(doc, children, level + 1, counters)

        # Also handle tech_requirements rendering within a section
        # (handled at engine level)


def _reset_counters_for_level(counters: dict, level: int):
    """Reset counters for levels deeper than the given level."""
    if level <= 0:
        counters['h2'] = 0
    if level <= 1:
        counters['h3'] = 0


def _render_h1(doc, title: str, num: int):
    """Render a first-level heading: 一、Title."""
    full_title = f'{to_chinese_numeral(num)}、{title}'
    para = add_paragraph_with_style(
        doc, full_title,
        font_name=FONT_TITLE, font_size=FONT_SIZE_H1, bold=True,
    )
    set_paragraph_spacing(para, before=PARA_SPACING_H1_BEFORE,
                          after=PARA_SPACING_H1_AFTER)
    # Apply outline level for TOC
    para.paragraph_format.outline_level = 0  # Level 1 in TOC


def _render_h2(doc, title: str, num: int):
    """Render a second-level heading: （一）Title."""
    full_title = f'{to_bracketed_chinese(num)}{title}'
    para = add_paragraph_with_style(
        doc, full_title,
        font_name=FONT_TITLE, font_size=FONT_SIZE_H2, bold=True,
    )
    set_paragraph_spacing(para, before=PARA_SPACING_H2_BEFORE,
                          after=PARA_SPACING_H2_AFTER)
    para.paragraph_format.outline_level = 1  # Level 2 in TOC


def _render_h3(doc, title: str, num: int):
    """Render a third-level heading: 1. Title."""
    full_title = f'{num}. {title}'
    para = add_paragraph_with_style(
        doc, full_title,
        font_name=FONT_SUBTITLE, font_size=FONT_SIZE_H3, bold=True,
    )
    para.paragraph_format.outline_level = 2  # Level 3 in TOC


def _render_body_paragraphs(doc, content: str):
    """Split content by blank lines and render as body paragraphs.

    Supports markdown-style image placeholder lines: ![alt](path).
    Each paragraph gets 宋体 小四, 1.5x line-spacing, first-line indent 2 chars.
    """
    blocks = _split_into_blocks(content)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph()
        run = para.add_run(block)
        set_run_font(run, FONT_BODY, FONT_SIZE_BODY)
        set_paragraph_spacing(para, line_spacing=LINE_SPACING)
        set_first_line_indent(para, chars=2.0)


def _split_into_blocks(content: str) -> list[str]:
    """Split text content into paragraph blocks by double-newline."""
    if not content:
        return []
    # Normalise line endings
    text = content.replace('\r\n', '\n').replace('\r', '\n')
    # Split on one or more blank lines
    import re
    blocks = re.split(r'\n\s*\n', text)
    return [b.strip() for b in blocks if b.strip()]
