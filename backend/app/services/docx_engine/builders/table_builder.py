"""Table builder.

Supports:
  - Regular data tables with header row, body, borders, and captions.
  - Special 技术参数响应表 (Technical Parameter Response Table):
    3 columns (序号 | 招标要求 | 投标响应) with widths 10% | 45% | 45%.
    Deviations in the 投标响应 column are marked in red.

Formatting (TEMPLATES.md):
  - Header: 宋体五号(10.5pt)加粗, grey background #D9D9D9
  - Body: 宋体五号(10.5pt)
  - Borders: black solid 0.5pt all around
  - Caption above table: "表X-X 标题" in 宋体小五(9pt)加粗, centred
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from ..styles import (
    FONT_BODY,
    FONT_SIZE_TABLE,
    FONT_SIZE_TABLE_TITLE,
    COLOR_TABLE_HEADER_BG,
    COLOR_BORDER,
    COLOR_RED,
    set_run_font,
    set_cell_shading,
    set_table_borders,
    add_paragraph_with_style,
)


# =============================================================================
# Public API
# =============================================================================

def build_table(doc, headers: list[str], rows: list[list[str]],
                caption: str = '', col_widths: list[float] | None = None) -> object:
    """Build a standard data table with caption above it.

    Args:
        doc: python-docx Document.
        headers: Column header texts.
        rows: List of rows; each row is a list of cell strings.
        caption: Table caption shown above the table (e.g. "表1-1 项目清单").
        col_widths: Optional list of column widths in Cm.

    Returns:
        The created Table object.
    """
    # ---- Caption ----
    if caption:
        cap = add_paragraph_with_style(
            doc, caption,
            font_name=FONT_BODY, font_size=FONT_SIZE_TABLE_TITLE, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        cap.paragraph_format.space_after = Pt(4)

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.autofit = True
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Header row ----
    _format_header_row(table.rows[0], headers)

    # ---- Body rows ----
    for row_idx, row_data in enumerate(rows):
        _format_body_row(table.rows[row_idx + 1], row_data)

    # ---- Column widths ----
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < num_cols:
                    row.cells[idx].width = width

    # ---- Borders ----
    set_table_borders(table, size=4, color=COLOR_BORDER)

    return table


def build_tech_response_table(doc, items: list[dict], caption: str = '') -> object:
    """Build the special 技术参数响应表 (Technical Parameter Response Table).

    3 columns: 序号 (10%) | 招标要求 (45%) | 投标响应 (45%)
    Cells in 投标响应 that contain deviation markers are coloured red.

    Args:
        doc: python-docx Document.
        items: list of dicts, each with keys: index, requirement, response.
        caption: Optional table caption.

    Returns:
        The created Table object.
    """
    headers = ['序号', '招标要求', '投标响应']
    col_widths = [Cm(1.6), Cm(7.2), Cm(7.2)]  # Approx 10/45/45 of A4 text width

    # ---- Caption ----
    if caption:
        cap = add_paragraph_with_style(
            doc, caption,
            font_name=FONT_BODY, font_size=FONT_SIZE_TABLE_TITLE, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(items), cols=3)
    table.autofit = True
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Header ----
    _format_header_row(table.rows[0], headers)
    for idx, width in enumerate(col_widths):
        table.rows[0].cells[idx].width = width

    # ---- Body ----
    for row_idx, item in enumerate(items):
        row = table.rows[row_idx + 1]
        index_text = str(item.get('index', row_idx + 1))
        requirement_text = item.get('requirement', '')
        response_text = item.get('response', '')
        has_deviation = item.get('deviation', False)

        cells_data = [index_text, requirement_text, response_text]
        for col_idx, cell_text in enumerate(cells_data):
            cell = row.cells[col_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT

            run = para.add_run(cell_text)
            font_color = COLOR_RED if (col_idx == 2 and has_deviation) else None
            set_run_font(run, FONT_BODY, FONT_SIZE_TABLE, color=font_color)

            cell.width = col_widths[col_idx]

    # ---- Borders ----
    set_table_borders(table, size=4, color=COLOR_BORDER)

    return table


# =============================================================================
# Internal helpers
# =============================================================================

def _format_header_row(row, headers: list[str]):
    """Format a table header row."""
    for idx, text in enumerate(headers):
        cell = row.cells[idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_run_font(run, FONT_BODY, FONT_SIZE_TABLE, bold=True)
        set_cell_shading(cell, COLOR_TABLE_HEADER_BG)


def _format_body_row(row, row_data: list[str]):
    """Format a regular body row."""
    for idx, text in enumerate(row_data):
        cell = row.cells[idx]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(str(text))
        set_run_font(run, FONT_BODY, FONT_SIZE_TABLE)
