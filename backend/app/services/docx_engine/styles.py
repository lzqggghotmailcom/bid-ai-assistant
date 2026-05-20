"""Formatting constants and helper functions for DOCX generation.

All measurements and styles strictly follow TEMPLATES.md specs.
"""

from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# =============================================================================
# Page Setup
# =============================================================================
PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)

# =============================================================================
# Font Names
# =============================================================================
FONT_BODY = '宋体'
FONT_TITLE = '黑体'
FONT_SUBTITLE = '楷体'
FONT_FZXBS = '方正小标宋简体'
FONT_FZXBS_FALLBACK = '黑体'

# =============================================================================
# Font Sizes (Chinese standard equivalents)
# =============================================================================
FONT_SIZE_BODY = Pt(12)               # 小四
FONT_SIZE_H1 = Pt(16)                 # 三号
FONT_SIZE_H2 = Pt(14)                 # 四号
FONT_SIZE_H3 = Pt(12)                 # 小四
FONT_SIZE_COVER_TITLE = Pt(22)        # 二号
FONT_SIZE_COVER_SUBTITLE = Pt(18)     # 小二
FONT_SIZE_COVER_BOTTOM = Pt(16)       # 三号
FONT_SIZE_TABLE = Pt(10.5)            # 五号
FONT_SIZE_TABLE_TITLE = Pt(9)         # 小五
FONT_SIZE_HEADER_FOOTER = Pt(9)       # 小五

# =============================================================================
# Spacing
# =============================================================================
LINE_SPACING = 1.5
LINE_SPACING_SINGLE = 1.0
PARA_SPACING_H1_BEFORE = Pt(12)
PARA_SPACING_H1_AFTER = Pt(6)
PARA_SPACING_H2_BEFORE = Pt(6)
PARA_SPACING_H2_AFTER = Pt(3)

# =============================================================================
# Colors
# =============================================================================
COLOR_TABLE_HEADER_BG = 'D9D9D9'
COLOR_BORDER = '000000'
COLOR_WATERMARK = 'BFBFBF'
COLOR_RED = 'FF0000'

# =============================================================================
# Chinese Numbering
# =============================================================================
_CN_DIGITS = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']


def to_chinese_numeral(n: int) -> str:
    """Convert an integer to a Chinese numeral string.

    Examples:
        1 -> 一, 10 -> 十, 11 -> 十一, 20 -> 二十, 99 -> 九十九
    """
    if n <= 0:
        return '零'
    if n <= 10:
        return _CN_DIGITS[n]
    if n < 20:
        return '十' + _CN_DIGITS[n - 10]
    tens = n // 10
    ones = n % 10
    if ones == 0:
        return _CN_DIGITS[tens] + '十'
    return _CN_DIGITS[tens] + '十' + _CN_DIGITS[ones]


def to_bracketed_chinese(n: int) -> str:
    """Convert integer to fullwidth-parenthesised Chinese numeral: （一）, （二）, etc."""
    return '（' + to_chinese_numeral(n) + '）'


def _fld_char(run, fld_char_type: str):
    """Append a w:fldChar element to a run."""
    el = OxmlElement('w:fldChar')
    el.set(qn('w:fldCharType'), fld_char_type)
    run._r.append(el)


def _instr_text(run, text: str):
    """Append a w:instrText element to a run."""
    el = OxmlElement('w:instrText')
    el.set(qn('xml:space'), 'preserve')
    el.text = text
    run._r.append(el)


# =============================================================================
# Public Helper Functions
# =============================================================================

def set_cell_border(cell, **kwargs):
    """Set individual cell borders.

    Usage:
        set_cell_border(cell, top={"sz": 4, "val": "single", "color": "000000"})
    Border size is in eighths of a point: 4 = 0.5pt, 8 = 1pt.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}></w:tcBorders>'
    )
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            attrs = kwargs[edge]
            el = parse_xml(
                f'<w:{edge} {nsdecls("w")} '
                f'w:val="{attrs.get("val", "single")}" '
                f'w:sz="{attrs.get("sz", 4)}" '
                f'w:space="0" '
                f'w:color="{attrs.get("color", "000000")}"/>'
            )
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders(table, size: int = 4, color: str = "000000", val: str = "single"):
    """Set all borders on a table uniformly.

    Args:
        size: Border width in eighths of a point (4 = 0.5pt).
        color: Hex color string.
        val: Border style ('single', 'double', etc.).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_cell_shading(cell, color: str):
    """Set cell background (shading) color.

    Args:
        color: Hex color string without # (e.g. 'D9D9D9').
    """
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, font_name: str = FONT_BODY, font_size=None,
                 bold: bool = False, color: str | None = None):
    """Configure font properties on a Run object, including East-Asian font.

    Args:
        run: A docx.text.run.Run object.
        font_name: Font family name (sets both ASCII and East-Asian).
        font_size: A Pt(...) value.
        bold: Whether text is bold.
        color: Hex colour string for the font colour.
    """
    if font_size is None:
        font_size = FONT_SIZE_BODY
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # Set East-Asian font face
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=LINE_SPACING):
    """Configure paragraph spacing.

    Args:
        before: Space before paragraph (Pt value or 0).
        after: Space after paragraph (Pt value or 0).
        line_spacing: Line spacing multiplier (e.g. 1.5).
    """
    pf = paragraph.paragraph_format
    if before:
        pf.space_before = before
    if after:
        pf.space_after = after
    pf.line_spacing = line_spacing


def set_first_line_indent(paragraph, chars: float = 2.0):
    """Set first-line indent in approximate character units.

    At 小四 (12pt), 2 chars = 24pt.
    """
    paragraph.paragraph_format.first_line_indent = int(FONT_SIZE_BODY * chars)


def add_paragraph_with_style(doc, text: str, font_name: str = FONT_BODY,
                             font_size=None, bold: bool = False,
                             alignment=None, color: str | None = None):
    """Add a paragraph to the document with full style control.

    Returns:
        The newly created Paragraph object.
    """
    if font_size is None:
        font_size = FONT_SIZE_BODY
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    if text:
        run = para.add_run(text)
        set_run_font(run, font_name, font_size, bold, color)
    return para


def add_field(paragraph, field_code: str, display_text: str = "1",
              font_name: str | None = None, font_size=None):
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC) into a paragraph.

    Args:
        field_code: The field instruction text (e.g. ' PAGE ', ' NUMPAGES ').
        display_text: Text shown as the field result placeholder.
        font_name: Optional font for the display-text run.
        font_size: Optional font size for the display-text run.
    """
    # Begin + instrText + separate
    run_begin = paragraph.add_run()
    _fld_char(run_begin, 'begin')
    _instr_text(run_begin, field_code)
    _fld_char(run_begin, 'separate')

    # Display text + end
    run_end = paragraph.add_run(display_text)
    if font_name is not None:
        set_run_font(run_end, font_name, font_size or FONT_SIZE_BODY)
    _fld_char(run_end, 'end')
