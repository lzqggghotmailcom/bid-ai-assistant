"""Main DOCX document generator.

Orchestrates all builders to produce a complete bid document:
  Cover -> TOC -> Content (sections + tables) -> Signature -> Watermark

Exports the single public function:
  generate_docx(outline_data, bid_info, output_path) -> output_path
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styles import (
    MARGIN_TOP,
    MARGIN_BOTTOM,
    MARGIN_LEFT,
    MARGIN_RIGHT,
    FONT_BODY,
    FONT_SIZE_HEADER_FOOTER,
    set_run_font,
    set_table_borders,
    add_field,
)
from .builders.cover_builder import build_cover
from .builders.toc_builder import build_toc
from .builders.content_builder import build_content
from .builders.table_builder import build_tech_response_table
from .builders.signature_builder import build_signature
from .watermark import add_watermark


def generate_docx(outline_data: dict, bid_info: dict, output_path: str) -> str:
    """Generate the complete bid Word document.

    Args:
        outline_data: dict with keys:
            sections (list[dict]) - outline sections with optional children
            tech_requirements (list[dict]|None) - for 技术参数响应表
        bid_info: dict with keys:
            company_name, project_name, bid_date, contact (optional)
        output_path: File-system path where the .docx will be written.

    Returns:
        The output_path string on success.
    """
    doc = Document()

    # ---- 1. Page Setup ----
    _apply_page_setup(doc)

    # ---- 2. Cover Page ----
    build_cover(doc, bid_info)
    doc.add_page_break()

    # ---- 3. Table of Contents ----
    build_toc(doc)
    doc.add_page_break()

    # ---- 4. Body Content ----
    sections = outline_data.get('sections', [])
    build_content(doc, sections)

    # ---- 5. Technical Parameter Response Table (if present) ----
    tech_reqs = outline_data.get('tech_requirements')
    if tech_reqs:
        doc.add_page_break()
        _build_tech_section(doc, tech_reqs)

    # ---- 6. Signature Page ----
    doc.add_page_break()
    build_signature(doc, bid_info)

    # ---- 7. Headers and Footers ----
    _add_header(doc, bid_info)
    _add_footer(doc)

    # ---- 8. Watermark ----
    add_watermark(doc)

    # ---- 9. Save ----
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_path


# =============================================================================
# Internal helpers
# =============================================================================

def _apply_page_setup(doc: Document):
    """Set A4 page size and standard margins on the default section."""
    section = doc.sections[0]
    section.page_width = doc.sections[0].page_width   # A4 is default
    section.page_height = doc.sections[0].page_height
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


def _add_header(doc: Document, bid_info: dict):
    """Add page header: company name (left) + 投标文件 (right), 宋体小五."""
    company = bid_info.get('company_name', '')
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Use a two-column table for left/right alignment
    hdr_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    hdr_table.autofit = True

    # Remove table borders (borderless table for layout)
    set_table_borders(hdr_table, size=0, color='FFFFFF', val='none')

    # Left cell: company name
    left_cell = hdr_table.rows[0].cells[0]
    left_cell.text = ''
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run(company)
    set_run_font(left_run, FONT_BODY, FONT_SIZE_HEADER_FOOTER)

    # Right cell: 投标文件
    right_cell = hdr_table.rows[0].cells[1]
    right_cell.text = ''
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run('投标文件')
    set_run_font(right_run, FONT_BODY, FONT_SIZE_HEADER_FOOTER)


def _add_footer(doc: Document):
    """Add page footer: centred page numbers in format '第X页 共Y页', 宋体小五."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.0

    # Build: 第 [PAGE] 页 共 [NUMPAGES] 页
    prefix_run = para.add_run('第')
    set_run_font(prefix_run, FONT_BODY, FONT_SIZE_HEADER_FOOTER)

    add_field(para, ' PAGE ', '1',
              font_name=FONT_BODY, font_size=FONT_SIZE_HEADER_FOOTER)

    mid_run = para.add_run('页 共')
    set_run_font(mid_run, FONT_BODY, FONT_SIZE_HEADER_FOOTER)

    add_field(para, ' NUMPAGES ', '1',
              font_name=FONT_BODY, font_size=FONT_SIZE_HEADER_FOOTER)

    suffix_run = para.add_run('页')
    set_run_font(suffix_run, FONT_BODY, FONT_SIZE_HEADER_FOOTER)


def _build_tech_section(doc, tech_items: list[dict]):
    """Render the technical parameter response table section."""
    caption = '技术参数响应表'
    build_tech_response_table(doc, tech_items, caption=caption)
