"""
Document-to-text converter.

Extracts the full plain-text content from Microsoft Word documents
(.docx, .doc) for feeding into the AI parsing pipeline.

* DOCX is handled natively via **python-docx**.
* DOC is handled via a LibreOffice subprocess call (headless mode).
  Falls back to a clear error when LibreOffice is not installed.
"""

from __future__ import annotations

import os
import subprocess
import tempfile
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> str:
    """
    Extract the full plain-text content of a document.

    Parameters
    ----------
    file_path : str
        Absolute or relative path to the document.

    Returns
    -------
    str
        Extracted text content.  Paragraphs are separated by double newlines.

    Raises
    ------
    FileNotFoundError
        If *file_path* does not exist.
    ValueError
        If the file extension is not supported.
    RuntimeError
        If conversion fails (e.g. corrupt file or missing LibreOffice).
    """
    path = Path(file_path).resolve()

    if not path.is_file():
        raise FileNotFoundError(f"Document not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx(path)
    elif suffix == ".doc":
        return _extract_doc(path)
    elif suffix == ".pdf":
        return _extract_pdf(path)
    else:
        raise ValueError(
            f"Unsupported document type: {suffix}. "
            f"Supported: .docx, .doc, .pdf"
        )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX extraction. "
            "Install it with: pip install python-docx"
        )

    logger.info("Extracting text from DOCX: %s", path)

    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        raise RuntimeError(f"File is not a valid DOCX (or is corrupt): {path}")
    except Exception as exc:
        raise RuntimeError(f"Failed to open DOCX: {path} — {exc}")

    parts: List[str] = []

    # --- Paragraphs ---------------------------------------------------------
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # --- Tables -------------------------------------------------------------
    for table in doc.tables:
        table_text = _extract_table_text(table)
        if table_text:
            parts.append(table_text)

    # --- Headers & Footers --------------------------------------------------
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    content = "\n\n".join(parts)
    logger.info("DOCX extracted: %d chars, %d paragraphs", len(content), len(parts))
    return content


def _extract_table_text(table) -> str:
    """Convert a python-docx Table to a pipe-delimited text block."""
    rows: List[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


# ---------------------------------------------------------------------------
# DOC (legacy binary format)
# ---------------------------------------------------------------------------

def _extract_doc(path: Path) -> str:
    """
    Extract text from a legacy .doc file via LibreOffice.

    This spawns a LibreOffice subprocess in headless mode to convert
    the .doc to .docx, then delegates to ``_extract_docx``.  The
    intermediate .docx is written to a temporary directory that is
    cleaned up automatically.

    If LibreOffice is not available, a RuntimeError is raised with
    installation instructions.
    """
    logger.info("Extracting text from DOC (via LibreOffice): %s", path)

    libreoffice = _find_libreoffice()
    if libreoffice is None:
        raise RuntimeError(
            "LibreOffice is required for legacy .doc conversion. "
            "Install it from https://www.libreoffice.org/ or via your "
            "package manager (apt install libreoffice-core, brew install "
            "libreoffice, etc.)."
        )

    with tempfile.TemporaryDirectory(prefix="bid_doc_convert_") as tmp_dir:
        cmd = [
            libreoffice,
            "--headless",
            "--convert-to", "docx",
            "--outdir", tmp_dir,
            str(path),
        ]

        logger.debug("Running: %s", " ".join(cmd))

        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,  # 2 minutes should be more than enough
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                "LibreOffice timed out converting .doc. The file may be "
                "too large or corrupt."
            )
        except FileNotFoundError:
            raise RuntimeError(f"LibreOffice executable not found: {libreoffice}")

        if result.returncode != 0:
            logger.error("LibreOffice stderr: %s", result.stderr)
            raise RuntimeError(
                f"LibreOffice conversion failed (exit code {result.returncode}). "
                f"stderr: {result.stderr[:500]}"
            )

        # Find the produced .docx — it will have the same stem.
        stem = path.stem
        converted = Path(tmp_dir) / f"{stem}.docx"
        if not converted.is_file():
            # LibreOffice may use a slightly different filename
            candidates = list(Path(tmp_dir).glob("*.docx"))
            if not candidates:
                raise RuntimeError(
                    "LibreOffice did not produce a .docx file. "
                    f"stdout: {result.stdout[:500]}"
                )
            converted = candidates[0]

        logger.info("DOC -> DOCX conversion successful, extracting text...")
        return _extract_docx(converted)


# ---------------------------------------------------------------------------
# PDF (via PyMuPDF) — lightweight fallback
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).

    This is a lightweight fallback for plain-text extraction.  The
    full-featured PDF parser with OCR support lives in
    ``app.services.pdf_parser``.
    """
    logger.info("Extracting text from PDF: %s", path)

    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError(
            "PyMuPDF is required for PDF extraction. "
            "Install it with: pip install PyMuPDF"
        )

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise RuntimeError(f"Failed to open PDF: {path} — {exc}")

    parts: List[str] = []
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            if text and text.strip():
                parts.append(text.strip())
    finally:
        doc.close()

    content = "\n\n".join(parts)
    logger.info("PDF extracted: %d chars, %d pages", len(content), len(parts))
    return content


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_libreoffice() -> str | None:
    """
    Locate the LibreOffice executable.

    Checks common paths on Linux, macOS, and Windows.
    """
    candidates = [
        # Linux
        "libreoffice",
        "soffice",
        # macOS
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        # Windows
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ]

    for candidate in candidates:
        if os.path.isfile(candidate) or shutil_which(candidate):
            return candidate

    return None


def shutil_which(cmd: str) -> str | None:
    """Vendored version of shutil.which for backwards compatibility."""
    import shutil
    return shutil.which(cmd)
