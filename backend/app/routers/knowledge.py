import os
import uuid

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.routers.auth import get_current_user
from app.core.config import settings
from app.core.database import get_db
from app.models.knowledge import KnowledgeDocument, KnowledgeChunk
from app.models.user import User
from app.schemas.knowledge import (
    KnowledgeDocResponse,
    KnowledgeListResponse,
    KnowledgeSearchRequest,
    KnowledgeSearchResponse,
    KnowledgeUploadResponse,
    SearchResultItem,
)
from app.services.rag.chunker import Chunker
from app.services.rag.embedder import Embedder
from app.services.rag.vector_store import insert_chunks, search as vector_search

router = APIRouter(prefix="/api/v1/knowledge", tags=["knowledge"])


# ---------------------------------------------------------------------------
# POST /knowledge/upload
# ---------------------------------------------------------------------------

@router.post("/upload", response_model=KnowledgeUploadResponse)
async def upload_knowledge(
    file: UploadFile = File(...),
    doc_type: str = Form(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload a file, parse it, chunk it, embed it, and store everything."""
    # --- Validate inputs ---
    if doc_type not in ("bid", "cert", "case", "resume"):
        raise HTTPException(
            status_code=400,
            detail=f"Invalid doc_type '{doc_type}'. Must be one of: bid, cert, case, resume.",
        )

    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")

    # --- Save file to disk ---
    upload_dir = os.path.join(settings.UPLOAD_DIR, str(current_user.id))
    os.makedirs(upload_dir, exist_ok=True)

    file_ext = os.path.splitext(file.filename)[1].lower()
    saved_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(upload_dir, saved_filename)

    content_bytes = await file.read()
    with open(file_path, "wb") as f:
        f.write(content_bytes)

    # --- Extract text ---
    try:
        text = _extract_text(file_path, file_ext)
    except Exception as exc:
        # Clean up the saved file on failure
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse file: {exc}",
        )

    if not text or not text.strip():
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=400,
            detail="No readable text found in the uploaded file.",
        )

    # --- Create document record ---
    doc = KnowledgeDocument(
        user_id=current_user.id,
        doc_type=doc_type,
        filename=file.filename,
        file_path=file_path,
        title=file.filename,
    )
    db.add(doc)
    await db.flush()  # populate doc.id without committing

    # --- Chunk ---
    chunker = Chunker()
    chunks = chunker.chunk(text)

    if not chunks:
        doc.status = "indexed"
        doc.chunk_count = 0
        await db.commit()
        return KnowledgeUploadResponse(
            doc_id=doc.id,
            chunks=0,
            status="indexed",
        )

    # --- Embed ---
    embedder = Embedder()
    chunk_texts = [c["content"] for c in chunks]
    embeddings = embedder.embed(chunk_texts)

    # --- Store chunks ---
    for chunk_dict in chunks:
        chunk_dict["document_id"] = doc.id
    await insert_chunks(db, chunks, embeddings)

    # --- Finalize document ---
    doc.chunk_count = len(chunks)
    doc.status = "indexed"
    await db.commit()

    return KnowledgeUploadResponse(
        doc_id=doc.id,
        chunks=len(chunks),
        status="indexed",
    )


# ---------------------------------------------------------------------------
# GET /knowledge/documents
# ---------------------------------------------------------------------------

@router.get("/documents", response_model=KnowledgeListResponse)
async def list_documents(
    doc_type: str | None = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List the current user's knowledge documents with pagination."""
    base_where = KnowledgeDocument.user_id == current_user.id

    # Count
    count_query = select(func.count(KnowledgeDocument.id)).where(base_where)
    if doc_type:
        count_query = count_query.where(KnowledgeDocument.doc_type == doc_type)
    total = (await db.execute(count_query)).scalar() or 0

    # Fetch page
    query = select(KnowledgeDocument).where(base_where)
    if doc_type:
        query = query.where(KnowledgeDocument.doc_type == doc_type)
    query = query.order_by(KnowledgeDocument.created_at.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(query)
    rows = result.scalars().all()

    return KnowledgeListResponse(
        items=[KnowledgeDocResponse.model_validate(doc) for doc in rows],
        total=total,
        page=page,
        page_size=page_size,
    )


# ---------------------------------------------------------------------------
# DELETE /knowledge/documents/{doc_id}
# ---------------------------------------------------------------------------

@router.delete("/documents/{doc_id}", status_code=204)
async def delete_document(
    doc_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Delete a knowledge document and all its chunks (cascading)."""
    result = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.id == doc_id,
            KnowledgeDocument.user_id == current_user.id,
        )
    )
    doc = result.scalar_one_or_none()
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found.")

    # Remove file from disk if it exists
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except OSError:
            pass  # best-effort cleanup

    await db.delete(doc)  # ON DELETE CASCADE removes chunks
    await db.commit()


# ---------------------------------------------------------------------------
# POST /knowledge/search
# ---------------------------------------------------------------------------

@router.post("/search", response_model=KnowledgeSearchResponse)
async def search_knowledge(
    req: KnowledgeSearchRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Semantic (vector) search over the user's knowledge base."""
    embedder = Embedder()
    query_embedding = embedder.embed([req.query])[0]

    results = await vector_search(
        db,
        query_embedding,
        current_user.id,
        top_k=req.top_k,
        doc_type=req.doc_type,
    )

    return KnowledgeSearchResponse(
        results=[
            SearchResultItem(
                doc_id=r["document_id"],
                content=r["content"],
                score=r["score"],
                section_title=r.get("section_title"),
            )
            for r in results
        ]
    )


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_text(file_path: str, file_ext: str) -> str:
    """Parse text content from a file based on its extension."""
    if file_ext in (".pdf",):
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        try:
            pages: list[str] = []
            for page in doc:
                pages.append(page.get_text())
            return "\n".join(pages)
        finally:
            doc.close()

    elif file_ext in (".docx", ".doc"):
        from docx import Document

        doc = Document(file_path)
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        return "\n".join(paragraphs)

    elif file_ext in (".txt", ".md", ".markdown"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    else:
        return ""
